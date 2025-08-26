from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def create_word_document(title, content, filename):
    """Create a Word document with proper formatting"""
    doc = Document()
    
    # Set up styles
    styles = doc.styles
    
    # Title style
    title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = 'Times New Roman'
    title_style.font.size = Pt(16)
    title_style.font.bold = True
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(12)
    
    # Heading 1 style
    h1_style = styles.add_style('CustomH1', WD_STYLE_TYPE.PARAGRAPH)
    h1_style.font.name = 'Times New Roman'
    h1_style.font.size = Pt(14)
    h1_style.font.bold = True
    h1_style.paragraph_format.space_before = Pt(12)
    h1_style.paragraph_format.space_after = Pt(6)
    
    # Heading 2 style
    h2_style = styles.add_style('CustomH2', WD_STYLE_TYPE.PARAGRAPH)
    h2_style.font.name = 'Times New Roman'
    h2_style.font.size = Pt(12)
    h2_style.font.bold = True
    h2_style.paragraph_format.space_before = Pt(12)
    h2_style.paragraph_format.space_after = Pt(6)
    
    # Heading 3 style
    h3_style = styles.add_style('CustomH3', WD_STYLE_TYPE.PARAGRAPH)
    h3_style.font.name = 'Times New Roman'
    h3_style.font.size = Pt(11)
    h3_style.font.bold = True
    h3_style.paragraph_format.space_before = Pt(6)
    h3_style.paragraph_format.space_after = Pt(3)
    
    # Normal text style
    normal_style = styles.add_style('CustomNormal', WD_STYLE_TYPE.PARAGRAPH)
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(12)
    normal_style.paragraph_format.space_after = Pt(6)
    normal_style.paragraph_format.first_line_indent = Inches(0.5)
    
    # List style
    list_style = styles.add_style('CustomList', WD_STYLE_TYPE.PARAGRAPH)
    list_style.font.name = 'Times New Roman'
    list_style.font.size = Pt(12)
    list_style.paragraph_format.space_after = Pt(3)
    list_style.paragraph_format.left_indent = Inches(0.5)
    
    # Add title
    title_para = doc.add_paragraph(title, style='CustomTitle')
    
    # Process content
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
            
        # Check for headings
        if line.startswith('# '):
            # Main title (already added)
            i += 1
            continue
        elif line.startswith('## '):
            # Heading 1
            heading = line[3:].strip()
            doc.add_paragraph(heading, style='CustomH1')
            i += 1
        elif line.startswith('### '):
            # Heading 2
            heading = line[4:].strip()
            doc.add_paragraph(heading, style='CustomH2')
            i += 1
        elif line.startswith('#### '):
            # Heading 3
            heading = line[5:].strip()
            doc.add_paragraph(heading, style='CustomH3')
            i += 1
        elif line.startswith('---'):
            # Horizontal line - skip
            i += 1
        elif line.startswith('```'):
            # Code block
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1  # Skip closing ```
            
            if code_lines:
                code_text = '\n'.join(code_lines)
                code_para = doc.add_paragraph(code_text, style='CustomNormal')
                code_para.style.font.name = 'Courier New'
                code_para.style.font.size = Pt(10)
        elif line.startswith('- ') or line.startswith('1. '):
            # List items
            # Remove markdown formatting and handle bold text
            clean_line = process_markdown_formatting(line)
            doc.add_paragraph(clean_line, style='CustomList')
            i += 1
        elif line.startswith('**') and line.endswith('**'):
            # Bold text only
            bold_text = line[2:-2]
            para = doc.add_paragraph(bold_text, style='CustomNormal')
            para.runs[0].bold = True
            i += 1
        elif line.startswith('```'):
            # Skip code block markers
            i += 1
        else:
            # Normal paragraph - process markdown formatting
            if line:
                clean_line = process_markdown_formatting(line)
                doc.add_paragraph(clean_line, style='CustomNormal')
            i += 1
    
    # Save document
    doc.save(filename)
    print(f"✅ Document saved as: {filename}")

def process_markdown_formatting(text):
    """Process markdown formatting and convert to Word formatting"""
    # Handle bold text **text** -> bold
    def bold_replace(match):
        return f"BOLD_START{match.group(1)}BOLD_END"
    
    text = re.sub(r'\*\*(.*?)\*\*', bold_replace, text)
    
    # Handle other markdown formatting
    text = text.replace('`', '')  # Remove backticks
    text = text.replace('*', '')  # Remove single asterisks
    
    return text

def add_formatted_text_to_paragraph(paragraph, text):
    """Add text with proper formatting to paragraph"""
    # Split by BOLD markers
    parts = text.split('BOLD_START')
    
    for i, part in enumerate(parts):
        if 'BOLD_END' in part:
            # This part contains bold text
            bold_text, remaining = part.split('BOLD_END', 1)
            run = paragraph.add_run(bold_text)
            run.bold = True
            run.font.name = 'Times New Roman'
            
            if remaining:
                run = paragraph.add_run(remaining)
                run.font.name = 'Times New Roman'
        else:
            # Regular text
            if part:
                run = paragraph.add_run(part)
                run.font.name = 'Times New Roman'

def create_word_document_improved(title, content, filename):
    """Create a Word document with improved formatting"""
    doc = Document()
    
    # Set up styles
    styles = doc.styles
    
    # Title style
    title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = 'Times New Roman'
    title_style.font.size = Pt(16)
    title_style.font.bold = True
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(12)
    
    # Heading 1 style
    h1_style = styles.add_style('CustomH1', WD_STYLE_TYPE.PARAGRAPH)
    h1_style.font.name = 'Times New Roman'
    h1_style.font.size = Pt(14)
    h1_style.font.bold = True
    h1_style.paragraph_format.space_before = Pt(12)
    h1_style.paragraph_format.space_after = Pt(6)
    
    # Heading 2 style
    h2_style = styles.add_style('CustomH2', WD_STYLE_TYPE.PARAGRAPH)
    h2_style.font.name = 'Times New Roman'
    h2_style.font.size = Pt(12)
    h2_style.font.bold = True
    h2_style.paragraph_format.space_before = Pt(12)
    h2_style.paragraph_format.space_after = Pt(6)
    
    # Heading 3 style
    h3_style = styles.add_style('CustomH3', WD_STYLE_TYPE.PARAGRAPH)
    h3_style.font.name = 'Times New Roman'
    h3_style.font.size = Pt(11)
    h3_style.font.bold = True
    h3_style.paragraph_format.space_before = Pt(6)
    h3_style.paragraph_format.space_after = Pt(3)
    
    # Normal text style
    normal_style = styles.add_style('CustomNormal', WD_STYLE_TYPE.PARAGRAPH)
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(12)
    normal_style.paragraph_format.space_after = Pt(6)
    normal_style.paragraph_format.first_line_indent = Inches(0.5)
    
    # List style
    list_style = styles.add_style('CustomList', WD_STYLE_TYPE.PARAGRAPH)
    list_style.font.name = 'Times New Roman'
    list_style.font.size = Pt(12)
    list_style.paragraph_format.space_after = Pt(3)
    list_style.paragraph_format.left_indent = Inches(0.5)
    
    # Add title
    title_para = doc.add_paragraph(title, style='CustomTitle')
    
    # Process content
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
            
        # Check for headings
        if line.startswith('# '):
            # Main title (already added)
            i += 1
            continue
        elif line.startswith('## '):
            # Heading 1
            heading = line[3:].strip()
            doc.add_paragraph(heading, style='CustomH1')
            i += 1
        elif line.startswith('### '):
            # Heading 2
            heading = line[4:].strip()
            doc.add_paragraph(heading, style='CustomH2')
            i += 1
        elif line.startswith('#### '):
            # Heading 3
            heading = line[5:].strip()
            doc.add_paragraph(heading, style='CustomH3')
            i += 1
        elif line.startswith('---'):
            # Horizontal line - skip
            i += 1
        elif line.startswith('```'):
            # Code block
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1  # Skip closing ```
            
            if code_lines:
                code_text = '\n'.join(code_lines)
                code_para = doc.add_paragraph(code_text, style='CustomNormal')
                code_para.style.font.name = 'Courier New'
                code_para.style.font.size = Pt(10)
        elif line.startswith('- ') or line.startswith('1. '):
            # List items
            clean_line = process_markdown_formatting(line)
            para = doc.add_paragraph(style='CustomList')
            add_formatted_text_to_paragraph(para, clean_line)
            i += 1
        elif line.startswith('```'):
            # Skip code block markers
            i += 1
        else:
            # Normal paragraph
            if line:
                clean_line = process_markdown_formatting(line)
                para = doc.add_paragraph(style='CustomNormal')
                add_formatted_text_to_paragraph(para, clean_line)
            i += 1
    
    # Save document
    doc.save(filename)
    print(f"✅ Document saved as: {filename}")

def convert_bab_ii():
    """Convert BAB II to Word format"""
    print("📝 Converting BAB II to Word format...")
    
    # Read BAB II content
    with open('vrp_rl_project/BAB_II_TINJAUAN_PUSTAKA.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create Word document
    create_word_document_improved(
        "BAB II TINJAUAN PUSTAKA",
        content,
        "vrp_rl_project/BAB_II_TINJAUAN_PUSTAKA_FIXED.docx"
    )

def convert_bab_iii():
    """Convert BAB III to Word format"""
    print("📝 Converting BAB III to Word format...")
    
    # Read BAB III content
    with open('vrp_rl_project/BAB_III_METODOLOGI_PENELITIAN.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create Word document
    create_word_document_improved(
        "BAB III METODOLOGI PENELITIAN",
        content,
        "vrp_rl_project/BAB_III_METODOLOGI_PENELITIAN_FIXED.docx"
    )

def convert_bab_iv():
    """Convert BAB IV to Word format"""
    print("📝 Converting BAB IV to Word format...")
    
    # Read BAB IV content
    with open('vrp_rl_project/BAB_IV_HASIL_DAN_PEMBAHASAN.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create Word document
    create_word_document_improved(
        "BAB IV HASIL DAN PEMBAHASAN",
        content,
        "vrp_rl_project/BAB_IV_HASIL_DAN_PEMBAHASAN.docx"
    )

def main():
    print("🔄 CONVERTING MARKDOWN TO WORD FORMAT (IMPROVED)")
    print("=" * 60)
    
    try:
        convert_bab_ii()
        convert_bab_iii()
        convert_bab_iv()
        
        print("\n🎉 CONVERSION COMPLETE!")
        print("=" * 60)
        print("📁 Files created:")
        print("   - BAB_II_TINJAUAN_PUSTAKA_FIXED.docx")
        print("   - BAB_III_METODOLOGI_PENELITIAN_FIXED.docx")
        print("   - BAB_IV_HASIL_DAN_PEMBAHASAN.docx")
        print("\n✨ Improvements:")
        print("   ✅ Bold text properly formatted")
        print("   ✅ Times New Roman font throughout")
        print("   ✅ Proper list formatting")
        print("   ✅ Clean markdown conversion")
        print("\n📋 Next Steps:")
        print("   1. Open the _FIXED.docx files in Microsoft Word")
        print("   2. Review formatting - should be much better now!")
        print("   3. Add page numbers, headers, footers as required")
        print("   4. Save as final version for your skripsi")
        
    except Exception as e:
        print(f"❌ Error during conversion: {e}")

if __name__ == "__main__":
    main() 